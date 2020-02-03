import sys
import os
import xlwt
from docx import Document

class Docx:
    def __init__(self, filename):
        self.filename = filename
    
    def open(self):
        self.document = Document(docx=os.path.join(os.getcwd(), 'default.docx'))
        # self.document = Document()

    def add_header(self, string):
        self.document.add_heading(string, level=1)
        self.table = self.document.add_table(rows=1, cols=5)
        self.table.style = 'Table Grid'
        hdr_cells = self.table.rows[0].cells
        hdr_cells[0].text = 'address'
        hdr_cells[1].text = 'parameter'
        hdr_cells[2].text = 'bit'
        hdr_cells[3].text = 'init value'
        hdr_cells[4].text = 'description'

    def write(self, address, name, base_size, init_value, description):
        row_cells = self.table.add_row().cells
        row_cells[0].text = str('0x%x'%(address))
        row_cells[1].text = name
        row_cells[2].text = str('%d'%base_size)
        row_cells[3].text = str('0x%02x'%(init_value))
        row_cells[4].text = description
    
    def close(self):
        self.document.save(self.filename)

class Excel:
    def __init__(self, filename):
        self.filename = filename

    def open(self):
        self.document = xlwt.Workbook(encoding='utf-8', style_compression=0)
        borders = xlwt.Borders()
        borders.left = 1
        borders.right = 1
        borders.top = 1
        borders.bottom = 1
        borders.bottom_colour=0x3A

        ptn = xlwt.Pattern()
        ptn.pattern = xlwt.Pattern.SOLID_PATTERN
        ptn.pattern_fore_colour = 1

        font = xlwt.Font()
        font.bold = True

        self.style1 = xlwt.XFStyle()
        self.style1.pattern = ptn
        self.style1.borders = borders
        self.style1.font = font

        self.style2 = xlwt.XFStyle()
        self.style2.pattern = ptn
        self.style2.borders = borders
    
    def add_header(self, string):
        self.sheet = self.document.add_sheet(string.replace("[", "_").replace("]",""), cell_overwrite_ok=True)

        self.sheet.write(0, 0, 'address', self.style1)
        self.sheet.write(0, 1, 'parameter', self.style1)
        self.sheet.write(0, 2, 'bitwidth', self.style1)
        self.sheet.write(0, 3, 'initial value', self.style1)
        self.sheet.write(0, 4, 'description', self.style1)
        self.sheet_row = 1
    
    def write(self, address, name, base_size, init_value, description):
        self.sheet.write(self.sheet_row, 0, str('0x%x'%(address)), self.style2)
        self.sheet.write(self.sheet_row, 1, name, self.style2)
        self.sheet.write(self.sheet_row, 2, str('%d'%base_size), self.style2)
        self.sheet.write(self.sheet_row, 3, str('0x%02x'%(init_value)), self.style2)
        self.sheet.write(self.sheet_row, 4, description, self.style2)
        self.sheet_row = self.sheet_row+1
    
    def close(self):
        self.document.save(self.filename)

class Txt:
    def __init__(self, filename):
        self.filename = filename
    
    def open(self):
        self.file = open(self.filename, 'w')
    
    def add_header(self, string):
        self.file.write('TITLE: ' + string + '\n')
    
    def write(self, address, name, base_size, init_value, description):
        self.file.write('0x%x '%address)
        self.file.write(name+' ')
        self.file.write('%d '%base_size)
        self.file.write('0x%02x '%init_value)
        self.file.write(description)
        self.file.write('\n')
    
    def close(self):
        self.file.close()

class Sysout:
    def open(self):
        pass
    
    def add_header(self, string):
        print('TITLE: ' + string)
    
    def write(self, address, name, base_size, init_value, description):
        print('0x%x %s %d %02x %s'%(address, name, base_size, init_value, description))
    
    def close(self):
        pass

class DocFactory:
    @staticmethod
    def doc_generator(filename):
        if ".docx" in filename:
            return Docx(filename)
        elif "xls" in filename:
            return Excel(filename)
        elif ".txt" in filename:
            return Txt(filename)
        else:
            return Sysout()

if __name__ == "__main__":
    file = DocFactory.doc_generator("hh")
    file.open()
    file.add_header("hh")
    file.write(0x8000, 'test', 8, 0x2, 'des')
    file.write(0x8000, 'test', 8, 0x2, 'des')
    file.add_header("hh")
    file.write(0x8000, 'test', 8, 0x2, 'des')
    file.write(0x8000, 'test', 8, 0x2, 'des')
    file.close()

        


